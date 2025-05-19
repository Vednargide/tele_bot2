import os
import logging
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes

# Logging configuration
logging.basicConfig(format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)

# Get bot token from environment variables
TOKEN = os.getenv("TOKEN")

# Extract text with font metadata from PDF
def extract_text_with_fonts(pdf_path):
    doc_data = []  # Stores text with font metadata
    try:
        pdf = fitz.open(pdf_path)
        for page in pdf:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                for line in block["lines"]:
                    line_data = []
                    for span in line["spans"]:
                        line_data.append({
                            "text": span["text"],
                            "font": span["font"],
                            "size": span["size"]
                        })
                    doc_data.append(line_data)
        pdf.close()
    except Exception as e:
        logger.error(f"Error extracting PDF data: {e}")
    return doc_data

# Write text with fonts to DOCX
def write_to_docx(doc_data, output_path):
    try:
        doc = Document()
        for line in doc_data:
            paragraph = doc.add_paragraph()
            for segment in line:
                run = paragraph.add_run(segment["text"])
                run.font.name = segment["font"]
                run.font.size = Pt(segment["size"])
        doc.save(output_path)
    except Exception as e:
        logger.error(f"Error writing to DOCX: {e}")

# Telegram command: Start
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Hi! Send me a PDF file, and I'll extract its text and fonts as DOCX for you.")

# Telegram handler: File upload
async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.message.document:
        document = update.message.document
        file_name = document.file_name
        file_extension = file_name.split(".")[-1].lower()

        if file_extension != "pdf":
            await update.message.reply_text("❌ Unsupported file type. Please upload a PDF file.")
            return

        # Download the file
        file = await context.bot.get_file(document.file_id)
        input_path = f"input_{file_name}"
        output_path = f"converted_{file_name.replace('.pdf', '.docx')}"
        await file.download_to_drive(input_path)

        # Extract fonts and text, then write to DOCX
        doc_data = extract_text_with_fonts(input_path)
        if doc_data:
            write_to_docx(doc_data, output_path)
            await update.message.reply_document(document=open(output_path, "rb"))
            os.remove(input_path)
            os.remove(output_path)
        else:
            await update.message.reply_text("❌ Conversion failed. Could not extract data from the PDF.")
    else:
        await update.message.reply_text("Please send a PDF document.")

# Main function to run the bot
def main():
    # Initialize the bot application
    application = Application.builder().token(TOKEN).build()

    # Add command handlers
    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))

    # Run the bot
    print("🤖 Bot is running...")
    application.run_polling()

if __name__ == "__main__":
    main()
